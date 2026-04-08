"""DOCX export helpers for tailored resume output."""

from __future__ import annotations

from io import BytesIO


def is_heading_line(line: str) -> bool:
    """Infer whether a resume line should be exported as a heading."""
    stripped = line.strip()
    return bool(stripped) and stripped == stripped.upper() and len(stripped.split()) <= 5


def create_docx_bytes(resume_text: str) -> bytes:
    """Convert tailored resume text into a DOCX document."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX export requires the optional 'python-docx' package.") from exc

    document = Document()
    for raw_line in resume_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            document.add_paragraph("")
            continue

        stripped = line.strip()
        if stripped.startswith(("- ", "* ", "• ", "▪ ", "◦ ")):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif is_heading_line(stripped):
            document.add_paragraph(stripped, style="Heading 2")
        else:
            document.add_paragraph(stripped)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
